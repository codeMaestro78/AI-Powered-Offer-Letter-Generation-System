import pandas as pd
import PyPDF2
from typing import List, Dict, Any, Optional
import re
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
import tempfile
import os
import time
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Enhanced document processor for the AI Offer Letter Generator
    Handles PDF, TXT, DOCX, and CSV files with intelligent chunking
    """
    
    def __init__(self, config=None):
        """Initialize the document processor with configuration"""
        # Default configuration if not provided
        self.config = config or self._get_default_config()
        
        
        # Initialize text splitter with configuration
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=getattr(self.config, 'CHUNK_SIZE', 1000),
            chunk_overlap=getattr(self.config, 'CHUNK_OVERLAP', 200),
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len
        )
        
        # Document type patterns for intelligent processing
        self.document_patterns = {
            'leave_policy': ['leave policy', 'leave management', 'vacation policy'],
            'travel_policy': ['travel policy', 'travel guidelines', 'business travel'],
            'offer_letter': ['offer letter', 'employment offer', 'job offer'],
            'benefits': ['benefits', 'employee benefits', 'compensation benefits'],
            'compensation': ['compensation', 'salary structure', 'pay scale'],
            'company_policy': ['company policy', 'employee handbook', 'code of conduct']
        }

    def _get_default_config(self):
        """Get default configuration if not provided"""
        class DefaultConfig:
            CHUNK_SIZE = 1000
            CHUNK_OVERLAP = 200
            MAX_FILE_SIZE_MB = 10
        
        return DefaultConfig()

    def process_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process CSV file and return structured data
        
        Args:
            file_path (str): Path to the CSV file
            
        Returns:
            List[Dict[str, Any]]: List of records from CSV
        """
        try:
            # Read CSV with error handling
            df = pd.read_csv(file_path, encoding='utf-8')
            
            # Basic validation
            if df.empty:
                logger.warning(f"CSV file {file_path} is empty")
                return []
            
            # Clean column names
            df.columns = df.columns.str.strip()
            
            # Log info about the CSV
            logger.info(f"Processed CSV with {len(df)} records and columns: {list(df.columns)}")
            
            return df.to_dict('records')
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    df.columns = df.columns.str.strip()
                    logger.info(f"Successfully read CSV with {encoding} encoding")
                    return df.to_dict('records')
                except:
                    continue
            
            logger.error(f"Could not read CSV file {file_path} with any encoding")
            return []
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            return []

    def process_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF {file_path} is encrypted, attempting to decrypt")
                    pdf_reader.decrypt('')
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
                
                # Clean up the text
                text = self._clean_extracted_text(text)
                
                logger.info(f"Successfully extracted {len(text)} characters from PDF")
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return ""
        
        return text

    def process_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # Clean up the text
            text = self._clean_extracted_text(text)
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return ""

    def process_txt(self, file_path: str) -> str:
        """
        Process TXT file
        
        Args:
            file_path (str): Path to the TXT file
            
        Returns:
            str: File content
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        logger.info(f"Successfully read TXT with {encoding} encoding")
                        break
                except:
                    continue
            else:
                logger.error(f"Could not read TXT file {file_path} with any encoding")
                return ""
                
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return ""
        
        # Clean up the text
        content = self._clean_extracted_text(content)
        logger.info(f"Successfully processed TXT file with {len(content)} characters")
        
        return content

    def _clean_extracted_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and formatting issues
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove page markers if they exist
        text = re.sub(r'\n--- Page \d+ ---\n', '\n', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[^\w\s\n\r\t.,;:!?()[\]{}"\'`~@#$%^&*+=<>/-]', '', text)
        
        return text.strip()

    def intelligent_chunk_documents(self, documents: Dict[str, str]) -> List[Dict[str, Any]]:
        """
        Intelligently chunk documents based on their type and content with enhanced context preservation
        
        Args:
            documents (Dict[str, str]): Dictionary of document names and content
            
        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        all_chunks = []
        
        for doc_name, content in documents.items():
            if not content or not content.strip():
                logger.warning(f"Empty content for document: {doc_name}")
                continue
                
            try:
                # Clean the content
                cleaned_content = self._clean_extracted_text(content)
                
                # Identify document type with enhanced detection
                doc_type = self._identify_document_type(doc_name, cleaned_content)
                
                # Extract document summary for better context
                doc_summary = self._extract_document_summary(cleaned_content, doc_name)
                
                # Apply specialized chunking based on document type
                if doc_type == 'travel_policy':
                    chunks = self._chunk_travel_policy(cleaned_content, doc_name)
                elif doc_type == 'leave_policy':
                    chunks = self._chunk_leave_policy(cleaned_content, doc_name)
                elif doc_type == 'offer_letter':
                    chunks = self._chunk_offer_letter(cleaned_content, doc_name)
                elif doc_type == 'benefits':
                    chunks = self._chunk_benefits_document(cleaned_content, doc_name)
                elif doc_type == 'compensation':
                    chunks = self._chunk_compensation_document(cleaned_content, doc_name)
                else:
                    chunks = self._enhanced_default_chunking(cleaned_content, doc_name)
                
                # Enhance chunks with better metadata and context
                enhanced_chunks = self._enhance_chunks_with_context(chunks, doc_summary, doc_type, cleaned_content)
                
                # Add document-level metadata to all chunks
                for chunk in enhanced_chunks:
                    chunk["metadata"]["document_type"] = doc_type
                    chunk["metadata"]["document_name"] = doc_name
                    chunk["metadata"]["document_summary"] = doc_summary
                    chunk["metadata"]["processing_timestamp"] = time.time()
                    chunk["metadata"]["word_count"] = len(chunk["content"].split())
                    chunk["metadata"]["char_count"] = len(chunk["content"])
                
                all_chunks.extend(enhanced_chunks)
                logger.info(f"Successfully chunked {doc_name} into {len(enhanced_chunks)} pieces (type: {doc_type})")
                
            except Exception as e:
                logger.error(f"Error processing document {doc_name}: {str(e)}")
                # Create a fallback chunk for failed documents
                fallback_chunk = {
                    "content": cleaned_content[:2000] + "..." if len(cleaned_content) > 2000 else cleaned_content,
                    "metadata": {
                        "source": doc_name,
                        "section": "Full_Document",
                        "type": "general",
                        "document_type": "unknown",
                        "document_name": doc_name,
                        "processing_timestamp": time.time(),
                        "error": f"Processing failed: {str(e)}",
                        "importance": "low",
                        "word_count": len(cleaned_content.split()),
                        "char_count": len(cleaned_content)
                    }
                }
                all_chunks.append(fallback_chunk)
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        return all_chunks

    def _identify_document_type(self, doc_name: str, doc_text: str) -> str:
        """
        Identify the type of document based on name and content
        
        Args:
            doc_name (str): Name of the document
            doc_text (str): Content of the document
            
        Returns:
            str: Document type identifier
        """
        doc_text_lower = doc_text.lower()
        doc_name_lower = doc_name.lower()

        # Add 'travel' as a standalone keyword for travel policy
        if 'travel_policy' in self.document_patterns:
            if 'travel' not in self.document_patterns['travel_policy']:
                 self.document_patterns['travel_policy'].append('travel')

        for doc_type, patterns in self.document_patterns.items():
            for pattern in patterns:
                if pattern.lower() in doc_text_lower or pattern.lower() in doc_name_lower:
                    logger.info(f"Document '{doc_name}' identified as type: {doc_type}")
                    return doc_type
        
        logger.info(f"Document '{doc_name}' identified as type: general")
        return 'general'

    def _chunk_leave_policy(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Specialized chunking for leave policy documents
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []

        # Extract Band-wise Leave Entitlement Matrix
        leave_matrix_patterns = [
            r"🏷\s*2\.\s*Band-wise Leave Entitlement Matrix(.*?)(?=🧠\s*3\.|🧾\s*4\.|\Z)",
            r"Band-wise Leave Entitlement(.*?)(?=Types of Leave|Leave Application|\Z)",
            r"Leave Entitlement Matrix(.*?)(?=Leave Types|Application Process|\Z)"
        ]
        
        for pattern in leave_matrix_patterns:
            leave_matrix = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if leave_matrix:
                chunks.append({
                    "content": leave_matrix.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Leave Entitlement Matrix",
                        "type": "leave_entitlement",
                        "importance": "high"
                    }
                })
                break

        # Extract Types of Leave Explained
        types_patterns = [
            r"🧠\s*3\.\s*Types of Leave Explained(.*?)(?=🧾\s*4\.|🏢\s*5\.|\Z)",
            r"Types of Leave(.*?)(?=Leave Application|WFO Expectations|\Z)",
            r"Leave Types(.*?)(?=Application|Work From Office|\Z)"
        ]
        
        for pattern in types_patterns:
            types = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if types:
                chunks.append({
                    "content": types.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Types of Leave",
                        "type": "leave_types",
                        "importance": "high"
                    }
                })
                break

        # Extract WFO Expectations
        wfo_patterns = [
            r"🏢\s*5\.\s*Work From Office.*?Expectations.*?by Team(.*?)(?=🏡\s*6\.|📋\s*7\.|\Z)",
            r"Work From Office.*?Expectations(.*?)(?=WFH Infrastructure|Support|\Z)",
            r"WFO Expectations(.*?)(?=WFH|Remote Work|\Z)"
        ]
        
        for pattern in wfo_patterns:
            wfo = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if wfo:
                chunks.append({
                    "content": wfo.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "WFO Expectations",
                        "type": "wfo_policy",
                        "importance": "medium"
                    }
                })
                break

        # If no specific sections found, fall back to default chunking
        if not chunks:
            chunks = self._default_chunking(content, doc_name)
            # Update metadata for leave policy
            for chunk in chunks:
                chunk["metadata"]["type"] = "leave_policy"

        return chunks

    def _chunk_travel_policy(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Specialized chunking for travel policy documents
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []

        # Extract Travel Eligibility & Entitlements
        eligibility_patterns = [
            r"2\.\s*Travel Eligibility & Entitlements \(Band-wise\)(.*?)(?=📅\s*3\.|Booking Process|Reimbursements|\Z)",
            r"Travel Eligibility.*?Entitlements(.*?)(?=Booking Process|Reimbursements|\Z)"
        ]
        for pattern in eligibility_patterns:
            eligibility = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if eligibility:
                chunks.append({
                    "content": eligibility.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Travel Eligibility & Entitlements",
                        "type": "travel_eligibility",
                        "importance": "high"
                    }
                })
                break

        # Extract Travel Band Matrix
        travel_matrix_patterns = [
            r"2\.1\s*Travel Band Matrix(.*?)(?=📅\s*3\.|💰\s*4\.|\Z)",
            r"Travel Band Matrix(.*?)(?=Booking Process|Reimbursements|\Z)",
            r"Band Matrix(.*?)(?=Booking|Travel Class|\Z)"
        ]
        
        for pattern in travel_matrix_patterns:
            travel_matrix = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if travel_matrix:
                chunks.append({
                    "content": travel_matrix.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Travel Band Matrix",
                        "type": "travel_eligibility",
                        "importance": "high"
                    }
                })
                break

        # Extract Reimbursement Policy
        reimburse_patterns = [
            r"💰\s*4\.\s*Reimbursements(.*?)(?=\Z|📋\s*5\.)",
            r"Reimbursement.*?Policy(.*?)(?=\Z|Additional|Notes)",
            r"Reimbursements(.*?)(?=\Z|Important|Note)"
        ]
        
        for pattern in reimburse_patterns:
            reimbursement = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if reimbursement:
                chunks.append({
                    "content": reimbursement.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Reimbursement Policy",
                        "type": "travel_reimbursement",
                        "importance": "high"
                    }
                })
                break

        # Extract Booking Process if present
        booking_patterns = [
            r"📅\s*3\.\s*Booking Process(.*?)(?=💰\s*4\.|\Z)",
            r"Booking Process(.*?)(?=Reimbursements|\Z)"
        ]
        
        for pattern in booking_patterns:
            booking = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if booking:
                chunks.append({
                    "content": booking.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Booking Process",
                        "type": "travel_booking",
                        "importance": "medium"
                    }
                })
                break

        # If no specific sections found, fall back to default chunking
        if not chunks:
            chunks = self._default_chunking(content, doc_name)
            # Update metadata for travel policy
            for chunk in chunks:
                chunk["metadata"]["type"] = "travel_policy"

        return chunks

    def _chunk_offer_letter(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Specialized chunking for offer letter documents
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []

        # Extract Compensation Structure
        comp_patterns = [
            r"💰\s*Compensation Structure.*?(?=🏖|🎯|📋|\Z)",
            r"Compensation.*?Structure(.*?)(?=Benefits|Leave|\Z)",
            r"Salary.*?Structure(.*?)(?=Benefits|Perks|\Z)"
        ]
        
        for pattern in comp_patterns:
            compensation = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if compensation:
                chunks.append({
                    "content": compensation.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Compensation Structure",
                        "type": "compensation",
                        "importance": "high"
                    }
                })
                break

        # Extract Benefits section
        benefits_patterns = [
            r"🏖.*?Benefits(.*?)(?=📋|🎯|\Z)",
            r"Employee.*?Benefits(.*?)(?=Terms|Conditions|\Z)",
            r"Benefits.*?Package(.*?)(?=Terms|\Z)"
        ]
        
        for pattern in benefits_patterns:
            benefits = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if benefits:
                chunks.append({
                    "content": benefits.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": "Benefits Package",
                        "type": "benefits",
                        "importance": "high"
                    }
                })
                break

        # If no specific sections found, fall back to default chunking
        if not chunks:
            chunks = self._default_chunking(content, doc_name)
            # Update metadata for offer letter
            for chunk in chunks:
                chunk["metadata"]["type"] = "offer_letter"

        return chunks

    def _chunk_benefits_document(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Specialized chunking for benefits documents
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []
        
        # Look for common benefits sections
        benefit_sections = [
            ("Health Insurance", r"Health.*?Insurance(.*?)(?=Dental|Vision|Life|\Z)"),
            ("Dental Coverage", r"Dental.*?Coverage(.*?)(?=Vision|Life|Retirement|\Z)"),
            ("Retirement Benefits", r"Retirement.*?Benefits(.*?)(?=Stock|Bonus|\Z)"),
            ("Stock Options", r"Stock.*?Options(.*?)(?=Bonus|Vacation|\Z)"),
            ("Vacation Policy", r"Vacation.*?Policy(.*?)(?=Sick|Personal|\Z)")
        ]
        
        for section_name, pattern in benefit_sections:
            match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if match:
                chunks.append({
                    "content": match.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": section_name,
                        "type": "benefits",
                        "importance": "high"
                    }
                })
        
        # If no specific sections found, use default chunking
        if not chunks:
            chunks = self._default_chunking(content, doc_name)
            for chunk in chunks:
                chunk["metadata"]["type"] = "benefits"
        
        return chunks

    def _chunk_compensation_document(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Specialized chunking for compensation documents
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []
        
        # Look for compensation-related sections
        comp_sections = [
            ("Salary Bands", r"Salary.*?Bands(.*?)(?=Bonus|Incentive|\Z)"),
            ("Bonus Structure", r"Bonus.*?Structure(.*?)(?=Stock|Equity|\Z)"),
            ("Equity Compensation", r"Equity.*?Compensation(.*?)(?=Benefits|\Z)"),
            ("Performance Incentives", r"Performance.*?Incentives(.*?)(?=Benefits|\Z)")
        ]
        
        for section_name, pattern in comp_sections:
            match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
            if match:
                chunks.append({
                    "content": match.group(0).strip(),
                    "metadata": {
                        "source": doc_name,
                        "section": section_name,
                        "type": "compensation",
                        "importance": "high"
                    }
                })
        
        # If no specific sections found, use default chunking
        if not chunks:
            chunks = self._default_chunking(content, doc_name)
            for chunk in chunks:
                chunk["metadata"]["type"] = "compensation"
        
        return chunks

    def _default_chunking(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Default chunking strategy for general documents
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []
        
        try:
            text_chunks = self.text_splitter.split_text(content)
            
            for i, chunk in enumerate(text_chunks):
                if chunk.strip():  # Only add non-empty chunks
                    chunks.append({
                        "content": chunk.strip(),
                        "metadata": {
                            "source": doc_name,
                            "section": f"Section_{i+1}",
                            "type": "general",
                            "chunk_index": i,
                            "total_chunks": len(text_chunks),
                            "importance": "medium"
                        }
                    })
                    
        except Exception as e:
            logger.error(f"Error in default chunking for {doc_name}: {str(e)}")
            # Create a single chunk if splitting fails
            chunks.append({
                "content": content[:2000] + "..." if len(content) > 2000 else content,
                "metadata": {
                    "source": doc_name,
                    "section": "Full_Document",
                    "type": "general",
                    "chunk_index": 0,
                    "total_chunks": 1,
                    "importance": "medium",
                    "error": "Chunking failed, truncated content"
                }
            })
        
        return chunks

    def _extract_document_summary(self, content: str, doc_name: str) -> str:
        """
        Extract a brief summary of the document for better context
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            str: Document summary
        """
        try:
            # Extract first few sentences as summary
            sentences = re.split(r'[.!?]+', content[:1000])
            summary_sentences = [s.strip() for s in sentences[:3] if s.strip() and len(s.strip()) > 10]
            
            if summary_sentences:
                summary = '. '.join(summary_sentences) + '.'
            else:
                # Fallback to first paragraph
                paragraphs = content.split('\n\n')
                summary = paragraphs[0][:200] + '...' if len(paragraphs[0]) > 200 else paragraphs[0]
            
            return summary.strip()
        except Exception as e:
            logger.warning(f"Could not extract summary for {doc_name}: {str(e)}")
            return f"Document: {doc_name}"

    def _enhance_chunks_with_context(self, chunks: List[Dict[str, Any]], doc_summary: str, 
                                   doc_type: str, full_content: str) -> List[Dict[str, Any]]:
        """
        Enhance chunks with additional context and improved metadata
        
        Args:
            chunks (List[Dict[str, Any]]): Original chunks
            doc_summary (str): Document summary
            doc_type (str): Document type
            full_content (str): Full document content
            
        Returns:
            List[Dict[str, Any]]: Enhanced chunks
        """
        enhanced_chunks = []
        
        for i, chunk in enumerate(chunks):
            enhanced_chunk = chunk.copy()
            
            # Add contextual information
            enhanced_chunk["metadata"]["chunk_position"] = i
            enhanced_chunk["metadata"]["total_chunks"] = len(chunks)
            enhanced_chunk["metadata"]["document_summary"] = doc_summary
            
            # Add surrounding context for better retrieval
            content = chunk["content"]
            
            # Find position in full document
            content_start = full_content.find(content[:100]) if len(content) > 100 else full_content.find(content)
            
            if content_start != -1:
                # Add preceding context (up to 200 chars)
                preceding_start = max(0, content_start - 200)
                preceding_context = full_content[preceding_start:content_start].strip()
                
                # Add following context (up to 200 chars)
                content_end = content_start + len(content)
                following_end = min(len(full_content), content_end + 200)
                following_context = full_content[content_end:following_end].strip()
                
                if preceding_context:
                    enhanced_chunk["metadata"]["preceding_context"] = preceding_context
                if following_context:
                    enhanced_chunk["metadata"]["following_context"] = following_context
            
            # Extract key terms for better searchability
            key_terms = self._extract_key_terms(content)
            enhanced_chunk["metadata"]["key_terms"] = key_terms
            
            # Determine chunk importance based on content
            importance = self._calculate_chunk_importance(content, doc_type)
            enhanced_chunk["metadata"]["importance"] = importance
            
            enhanced_chunks.append(enhanced_chunk)
        
        return enhanced_chunks

    def _extract_key_terms(self, content: str) -> List[str]:
        """
        Extract key terms from content for better searchability
        
        Args:
            content (str): Chunk content
            
        Returns:
            List[str]: Key terms
        """
        try:
            # Extract important terms (capitalized words, numbers with units, etc.)
            key_terms = []
            
            # Capitalized words (likely important terms)
            capitalized = re.findall(r'\b[A-Z][a-z]+\b', content)
            key_terms.extend(capitalized[:10])  # Limit to top 10
            
            # Numbers with units (salary, days, percentages)
            numbers_with_units = re.findall(r'\b\d+[%₹$,\s]*(?:days?|months?|years?|INR|USD|percent|%)\b', content, re.IGNORECASE)
            key_terms.extend(numbers_with_units[:5])
            
            # Policy-specific terms
            policy_terms = re.findall(r'\b(?:policy|procedure|guideline|requirement|benefit|entitlement|allowance)\b', content, re.IGNORECASE)
            key_terms.extend(policy_terms[:5])
            
            return list(set(key_terms))  # Remove duplicates
        except Exception as e:
            logger.warning(f"Could not extract key terms: {str(e)}")
            return []

    def _calculate_chunk_importance(self, content: str, doc_type: str) -> str:
        """
        Calculate the importance of a chunk based on its content and document type
        
        Args:
            content (str): Chunk content
            doc_type (str): Document type
            
        Returns:
            str: Importance level (high, medium, low)
        """
        try:
            content_lower = content.lower()
            importance_score = 0
            
            # High importance indicators
            high_importance_terms = [
                'salary', 'compensation', 'ctc', 'benefits', 'policy', 'requirement',
                'mandatory', 'must', 'shall', 'entitled', 'eligible', 'band', 'grade',
                'allowance', 'reimbursement', 'approval', 'process', 'procedure'
            ]
            
            for term in high_importance_terms:
                if term in content_lower:
                    importance_score += 2
            
            # Medium importance indicators
            medium_importance_terms = [
                'guideline', 'recommendation', 'should', 'may', 'can', 'option',
                'additional', 'supplementary', 'example', 'note'
            ]
            
            for term in medium_importance_terms:
                if term in content_lower:
                    importance_score += 1
            
            # Document type specific scoring
            if doc_type in ['travel_policy', 'leave_policy', 'compensation']:
                importance_score += 1
            
            # Length-based scoring (longer chunks often more important)
            if len(content) > 500:
                importance_score += 1
            
            # Determine final importance
            if importance_score >= 4:
                return 'high'
            elif importance_score >= 2:
                return 'medium'
            else:
                return 'low'
                
        except Exception as e:
            logger.warning(f"Could not calculate chunk importance: {str(e)}")
            return 'medium'

    def _enhanced_default_chunking(self, content: str, doc_name: str) -> List[Dict[str, Any]]:
        """
        Enhanced default chunking strategy with better context preservation
        
        Args:
            content (str): Document content
            doc_name (str): Document name
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []
        
        try:
            # Use semantic-aware splitting
            # First try to split by sections (headers, numbered items)
            section_patterns = [
                r'\n\s*(?:\d+\.\s+|[A-Z][^\n]*:)\s*\n',  # Numbered sections or headers
                r'\n\s*[A-Z][A-Z\s]{10,}\n',  # All caps headers
                r'\n\s*[-=]{3,}\s*\n',  # Separator lines
            ]
            
            sections = [content]  # Start with full content
            
            for pattern in section_patterns:
                new_sections = []
                for section in sections:
                    parts = re.split(pattern, section)
                    new_sections.extend([part.strip() for part in parts if part.strip()])
                sections = new_sections
            
            # If sections are too large, use text splitter
            final_chunks = []
            for section in sections:
                if len(section) > self.config.CHUNK_SIZE * 1.5:
                    # Use text splitter for large sections
                    text_chunks = self.text_splitter.split_text(section)
                    final_chunks.extend(text_chunks)
                else:
                    final_chunks.append(section)
            
            # Create chunk objects with enhanced metadata
            for i, chunk_content in enumerate(final_chunks):
                if chunk_content.strip():  # Only add non-empty chunks
                    # Try to identify section header
                    lines = chunk_content.strip().split('\n')
                    section_name = lines[0][:50] + '...' if len(lines[0]) > 50 else lines[0]
                    
                    chunks.append({
                        "content": chunk_content.strip(),
                        "metadata": {
                            "source": doc_name,
                            "section": section_name,
                            "type": "general",
                            "chunk_index": i,
                            "total_chunks": len(final_chunks),
                            "importance": "medium",
                            "processing_method": "enhanced_default"
                        }
                    })
                    
        except Exception as e:
            logger.error(f"Error in enhanced default chunking for {doc_name}: {str(e)}")
            # Fallback to simple chunking
            text_chunks = self.text_splitter.split_text(content)
            for i, chunk in enumerate(text_chunks):
                if chunk.strip():
                    chunks.append({
                        "content": chunk.strip(),
                        "metadata": {
                            "source": doc_name,
                            "section": f"Section_{i+1}",
                            "type": "general",
                            "chunk_index": i,
                            "total_chunks": len(text_chunks),
                            "importance": "medium",
                            "processing_method": "fallback",
                            "error": "Enhanced chunking failed"
                        }
                    })
        
        return chunks

    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Get statistics about document processing
        
        Returns:
            Dict[str, Any]: Processing statistics
        """
        return {
            "chunk_size": getattr(self.config, 'CHUNK_SIZE', 1000),
            "chunk_overlap": getattr(self.config, 'CHUNK_OVERLAP', 200),
            "supported_formats": ["PDF", "TXT", "DOCX", "CSV"],
            "document_types": list(self.document_patterns.keys()),
            "text_splitter_separators": self.text_splitter._separators
        }

    def validate_file_size(self, file_path: str, max_size_mb: Optional[int] = None) -> bool:
        """
        Validate file size before processing
        
        Args:
            file_path (str): Path to the file
            max_size_mb (int, optional): Maximum file size in MB
            
        Returns:
            bool: True if file size is acceptable
        """
        if max_size_mb is None:
            max_size_mb = getattr(self.config, 'MAX_FILE_SIZE_MB', 10)
        
        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > max_size_mb:
                logger.warning(f"File {file_path} exceeds size limit: {file_size_mb:.2f}MB > {max_size_mb}MB")
                return False
            return True
        except Exception as e:
            logger.error(f"Error checking file size for {file_path}: {str(e)}")
            return False